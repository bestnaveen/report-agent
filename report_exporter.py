import io
from datetime import datetime
import pandas as pd


def export_to_csv(df: pd.DataFrame) -> bytes:
    """Export DataFrame to CSV bytes."""
    buf = io.StringIO()
    df.to_csv(buf, index=False)
    return buf.getvalue().encode("utf-8")


def export_to_docx(
    title: str,
    content: str,
    df: pd.DataFrame | None = None,
    chart_images: list[bytes] | None = None,
) -> bytes:
    """
    Export report to DOCX bytes.

    Args:
        title: Report title
        content: Text/markdown content from AI
        df: Optional DataFrame to include as table (max 100 rows)
        chart_images: Optional list of PNG bytes to embed
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Timestamp
    ts_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_para.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # Content — parse markdown-ish text
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            # Markdown table row — skip for now, will render as DataFrame below
            pass
        else:
            doc.add_paragraph(line)

    # Chart images
    if chart_images:
        doc.add_heading("Charts", level=2)
        for i, img_bytes in enumerate(chart_images):
            if img_bytes:
                doc.add_paragraph(f"Chart {i + 1}:")
                buf = io.BytesIO(img_bytes)
                doc.add_picture(buf, width=Inches(6))
                doc.add_paragraph()

    # Data table
    if df is not None and not df.empty:
        doc.add_heading("Data Table", level=2)
        rows_to_show = df.head(100)
        table = doc.add_table(rows=1 + len(rows_to_show), cols=len(df.columns))
        table.style = "Table Grid"

        # Header row
        for j, col in enumerate(df.columns):
            cell = table.rows[0].cells[j]
            cell.text = str(col)
            run = cell.paragraphs[0].runs[0]
            run.bold = True

        # Data rows
        for i, (_, row) in enumerate(rows_to_show.iterrows()):
            for j, val in enumerate(row):
                table.rows[i + 1].cells[j].text = str(val) if pd.notna(val) else ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_to_pdf(
    title: str,
    content: str,
    df: pd.DataFrame | None = None,
    chart_images: list[bytes] | None = None,
) -> bytes:
    """
    Export report to PDF bytes using reportlab.

    Args:
        title: Report title
        content: Text content from AI
        df: Optional DataFrame (max 50 rows)
        chart_images: Optional list of PNG bytes to embed
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        fontSize=20,
        spaceAfter=6,
        textColor=colors.HexColor("#1a1a2e"),
    )
    h2_style = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=14,
        spaceBefore=12,
        textColor=colors.HexColor("#16213e"),
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=4,
    )
    ts_style = ParagraphStyle(
        "Timestamp",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.grey,
        spaceAfter=12,
    )

    story = []
    story.append(Paragraph(title, title_style))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ts_style
    ))
    story.append(Spacer(1, 0.2 * inch))

    # Content
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.1 * inch))
            continue
        if line.startswith("# "):
            story.append(Paragraph(line[2:], styles["Heading1"]))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], h2_style))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], styles["Heading3"]))
        elif line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph(f"• {line[2:]}", body_style))
        elif line.startswith("|"):
            pass  # skip markdown table lines; will render as PDF table below
        else:
            # Escape special reportlab chars
            safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line, body_style))

    # Chart images
    if chart_images:
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("Charts", h2_style))
        for img_bytes in chart_images:
            if img_bytes:
                img_buf = io.BytesIO(img_bytes)
                rl_img = RLImage(img_buf, width=6 * inch, height=3.3 * inch)
                story.append(rl_img)
                story.append(Spacer(1, 0.15 * inch))

    # DataFrame table
    if df is not None and not df.empty:
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("Data Table", h2_style))

        rows_to_show = df.head(50)
        table_data = [list(df.columns)]
        for _, row in rows_to_show.iterrows():
            table_data.append([str(v) if pd.notna(v) else "" for v in row])

        col_width = (A4[0] - 1.5 * inch) / len(df.columns)
        pdf_table = Table(table_data, colWidths=[col_width] * len(df.columns))
        pdf_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f0f0")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("PADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(pdf_table)

    doc.build(story)
    return buf.getvalue()
